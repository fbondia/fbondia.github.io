#!/usr/bin/env python3

from __future__ import annotations

import io
import json
import re
import subprocess
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLACK = "111111"
MUTED = "555555"
LIGHT_GRAY = "DEDEDE"
RULE = "555555"
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)
FONT = "Arial"


def load_yaml(path: Path):
    ruby = (
        "require 'yaml'; require 'json'; "
        "data = YAML.safe_load(File.read(ARGV[0]), permitted_classes: [Date, Time], aliases: true); "
        "puts JSON.generate(data)"
    )
    completed = subprocess.run(
        ["ruby", "-rdate", "-e", ruby, str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(completed.stdout)


def set_run_font(run, size=9.25, bold=None, italic=None, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border_paragraph(paragraph, color=RULE, size="5", position="bottom"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = borders.find(qn(f"w:{position}"))
    if border is None:
        border = OxmlElement(f"w:{position}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_hyperlink(paragraph, text, url, color=LINK_BLUE, underline=True):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    run_properties.append(size)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), "%02X%02X%02X" % tuple(color))
    run_properties.append(color_node)
    if underline:
        underline_node = OxmlElement("w:u")
        underline_node.set(qn("w:val"), "single")
        run_properties.append(underline_node)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clean_markup(text):
    text = re.sub(r"</?(?:i|em|b|strong)>", "", str(text), flags=re.I)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text.strip()


def split_paragraphs(text):
    return [clean_markup(block) for block in re.split(r"\n\s*\n", text or "") if clean_markup(block)]


def add_body_paragraph(doc, text, italic=False, after=3):
    paragraph = doc.add_paragraph(style="Resume Body")
    paragraph.paragraph_format.space_after = Pt(after)
    set_run_font(paragraph.add_run(clean_markup(text)), italic=italic)
    return paragraph


def add_bullet(doc, text, style="Resume Bullet"):
    paragraph = doc.add_paragraph(style=style)
    set_run_font(paragraph.add_run(clean_markup(text)))
    return paragraph


def add_section_band(doc, text):
    paragraph = doc.add_paragraph(style="Resume Section")
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(5)
    shade_paragraph(paragraph, LIGHT_GRAY)
    border_paragraph(paragraph, position="top")
    border_paragraph(paragraph, position="bottom")
    set_run_font(paragraph.add_run(text), size=10.5, bold=True)
    keep_with_next(paragraph)
    return paragraph


def add_rule(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    border_paragraph(paragraph, color="222222", size="5")


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])
    for item in paragraph.runs:
        set_run_font(item, size=8, color=MUTED)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.25)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    body = styles.add_style("Resume Body", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(3)
    body.paragraph_format.line_spacing = 1.05

    section = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    section.base_style = normal
    section.paragraph_format.keep_with_next = True

    role = styles.add_style("Resume Role", WD_STYLE_TYPE.PARAGRAPH)
    role.base_style = normal
    role.font.bold = True
    role.font.size = Pt(10.25)
    role.paragraph_format.space_after = Pt(0)
    role.paragraph_format.keep_with_next = True

    meta = styles.add_style("Resume Meta", WD_STYLE_TYPE.PARAGRAPH)
    meta.base_style = normal
    meta.font.size = Pt(9)
    meta.paragraph_format.space_after = Pt(3)
    meta.paragraph_format.keep_with_next = True

    category = styles.add_style("Resume Category", WD_STYLE_TYPE.PARAGRAPH)
    category.base_style = normal
    category.font.bold = True
    category.font.size = Pt(9.25)
    category.paragraph_format.space_before = Pt(3)
    category.paragraph_format.space_after = Pt(1)
    category.paragraph_format.keep_with_next = True

    bullet = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(9.25)
    bullet.paragraph_format.left_indent = Inches(0.24)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)
    bullet.paragraph_format.space_after = Pt(1)
    bullet.paragraph_format.line_spacing = 1.0

    compact_bullet = styles.add_style("Resume Compact Bullet", WD_STYLE_TYPE.PARAGRAPH)
    compact_bullet.base_style = bullet
    compact_bullet.font.size = Pt(8.75)
    compact_bullet.paragraph_format.space_after = Pt(0)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)


def set_columns(section, count=1, space=360):
    section_properties = section._sectPr
    columns = section_properties.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section_properties.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), str(space))


def add_column_break(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    break_node = OxmlElement("w:br")
    break_node.set(qn("w:type"), "column")
    run._r.append(break_node)


def add_header(document, profile, site_url):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    set_run_font(title.add_run(profile["name"]), size=18, bold=True)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    links = [
        (profile["email"], f"mailto:{profile['email']}"),
        (site_url.removeprefix("https://").removeprefix("http://"), site_url),
        (f"linkedin.com/in/{profile['linkedin']}", f"https://www.linkedin.com/in/{profile['linkedin']}/"),
        (f"github.com/{profile['github']}", f"https://github.com/{profile['github']}"),
    ]
    for index, (label, url) in enumerate(links):
        if index:
            set_run_font(contact.add_run("  ·  "), size=8.5)
        add_hyperlink(contact, label, url)

    objective = document.add_paragraph()
    objective.alignment = WD_ALIGN_PARAGRAPH.CENTER
    objective.paragraph_format.space_before = Pt(1)
    objective.paragraph_format.space_after = Pt(8)
    objective.paragraph_format.keep_with_next = True
    border_paragraph(objective, color="222222", size="6", position="top")
    border_paragraph(objective, color="222222", size="6", position="bottom")
    set_run_font(objective.add_run(profile["tagline"]), size=9.5, bold=True)


def add_about(document, about, lang):
    add_section_band(document, about["title"])
    summary = about["summary"]
    before_list = summary.split("<ul>", 1)[0]
    before_list = re.sub(r"<b><em>.*?</em></b>", "", before_list, flags=re.I | re.S)
    for block in split_paragraphs(before_list):
        add_body_paragraph(document, block)
    label = "Destaques" if lang == "pt" else "Highlights"
    paragraph = document.add_paragraph(style="Resume Category")
    set_run_font(paragraph.add_run(label), size=9.25, bold=True)
    for item in re.findall(r"<li>(.*?)</li>", summary, flags=re.I | re.S):
        add_bullet(document, item)


def add_languages_and_education(document, sidebar, education, lang):
    labels = {
        "languages": "Idiomas" if lang == "pt" else "Languages",
        "education": "Educação" if lang == "pt" else "Education",
    }
    add_section_band(document, labels["languages"])
    for item in sidebar["languages"]["list"]:
        add_bullet(document, f"{item['idiom']} ({item['level']})")

    add_section_band(document, labels["education"])
    for item in education:
        paragraph = document.add_paragraph(style="Resume Body")
        set_run_font(paragraph.add_run(item["degree"]), bold=False)
        set_run_font(paragraph.add_run(f"\n{item['university']}"), bold=True)
        set_run_font(paragraph.add_run(f"\n{item['time']}"), size=8.75, bold=True)
        paragraph.paragraph_format.space_after = Pt(4)


def add_skills(document, skills):
    add_section_band(document, skills["title"])
    section = document.add_section(WD_SECTION.CONTINUOUS)
    configure_section(section)
    set_columns(section, 2, 420)
    categories = skills["categories"]
    split_at = (len(categories) + 1) // 2
    for index, category in enumerate(categories):
        if index == split_at:
            add_column_break(document)
        heading = document.add_paragraph(style="Resume Category")
        set_run_font(heading.add_run(category["name"]), size=9, bold=True)
        for skill in category["toolset"]:
            add_bullet(document, skill["name"], style="Resume Compact Bullet")
    section = document.add_section(WD_SECTION.CONTINUOUS)
    configure_section(section)
    set_columns(section, 1)


def project_logo(path):
    image = Image.open(path).convert("RGBA")
    canvas = Image.new("RGBA", (120, 120), (255, 255, 255, 0))
    image.thumbnail((104, 104), Image.Resampling.LANCZOS)
    if path.name == "project-manura.png":
        canvas = Image.new("RGBA", (120, 120), (52, 52, 52, 255))
    x = (120 - image.width) // 2
    y = (120 - image.height) // 2
    canvas.alpha_composite(image, (x, y))
    stream = io.BytesIO()
    canvas.save(stream, format="PNG")
    stream.seek(0)
    return stream


def add_projects(document, projects, project_dir):
    add_section_band(document, projects["title"])
    for project in projects["items"]:
        header = document.add_paragraph(style="Resume Role")
        logo_path = project_dir / "assets" / "images" / project.get("logo", "")
        if logo_path.is_file():
            picture = header.add_run()
            picture.add_picture(project_logo(logo_path), width=Inches(0.29), height=Inches(0.29))
            set_run_font(header.add_run("  "), size=9)
        set_run_font(header.add_run(f"{project['name']} — {project['role']}"), size=10.25, bold=True)
        for block in split_paragraphs(project.get("description", "")):
            add_body_paragraph(document, block)
        stack = document.add_paragraph(style="Resume Body")
        set_run_font(stack.add_run("Stack: "), bold=True)
        set_run_font(stack.add_run(project["technologies"]))
        link = document.add_paragraph(style="Resume Body")
        add_hyperlink(link, project["link_label"], project["url"])
        add_rule(document)


def add_experiences(document, experiences):
    add_section_band(document, experiences["title"])
    for experience in experiences["roles"]:
        role = document.add_paragraph(style="Resume Role")
        set_run_font(role.add_run(experience["role"]), size=10.25, bold=True)
        company = document.add_paragraph(style="Resume Meta")
        set_run_font(company.add_run(f"{experience['company']} - {experience['time']}"), size=9)
        for block in split_paragraphs(experience.get("details", "")):
            italic = "<i>" in experience.get("details", "") and len(split_paragraphs(experience.get("details", ""))) <= 2
            add_body_paragraph(document, block, italic=italic)
        for highlight in experience.get("highlights", []):
            heading = document.add_paragraph(style="Resume Category")
            set_run_font(heading.add_run(highlight["title"]), size=9.1, bold=True)
            for item in highlight.get("items", []):
                add_bullet(document, item["name"], style="Resume Compact Bullet")
        add_rule(document)


def add_footer(section, lang):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_page_field(paragraph)


def build_resume(project_dir: Path, output_dir: Path, lang: str):
    config = load_yaml(project_dir / "_config.yml")
    sidebar = load_yaml(project_dir / "_data" / "sidebar.yml")[lang]
    about = load_yaml(project_dir / "_data" / "about.yml")[lang]
    education = load_yaml(project_dir / "_data" / "education.yml")[lang]
    skills = load_yaml(project_dir / "_data" / "skills.yml")[lang]
    projects = load_yaml(project_dir / "_data" / "projects.yml")[lang]
    experiences = load_yaml(project_dir / "_data" / "experience.yml")[lang]

    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])

    add_header(document, sidebar, config["url"])
    add_about(document, about, lang)
    add_languages_and_education(document, sidebar, education, lang)
    add_skills(document, skills)
    add_projects(document, projects, project_dir)
    add_experiences(document, experiences)

    for section in document.sections:
        configure_section(section)
    add_footer(document.sections[0], lang)

    document.core_properties.title = f"Fabiano Bondia - Resume ({lang.upper()})"
    document.core_properties.author = "Fabiano Bondia"
    document.core_properties.subject = "Professional Resume"
    document.core_properties.keywords = "resume, software engineering, software architecture, AI"

    output_path = output_dir / f"resume-{lang}.docx"
    document.save(output_path)
    print(f"Generated {output_path}")


def main():
    if len(sys.argv) != 4:
        raise SystemExit("Usage: generate-resume-docx.py [en|pt|all] OUTPUT_DIR PROJECT_DIR")
    language = sys.argv[1]
    output_dir = Path(sys.argv[2]).resolve()
    project_dir = Path(sys.argv[3]).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    languages = ("en", "pt") if language == "all" else (language,)
    for lang in languages:
        build_resume(project_dir, output_dir, lang)


if __name__ == "__main__":
    main()
